from io import BytesIO
from pathlib import Path
from functools import lru_cache
from threading import Lock

from docx import Document
from pypdf import PdfReader
import pypdfium2 as pdfium
from rapidocr import RapidOCR


SUPPORTED_EXTENSIONS = frozenset({".txt", ".md", ".docx", ".pdf"})
MAX_OCR_PAGES = 30
OCR_DPI = 160
_OCR_LOCK = Lock()


class DocumentParseError(ValueError):
    """上传文档格式不支持、内容损坏或无法提取文本。"""


@lru_cache(maxsize=1)
def _ocr_engine() -> RapidOCR:
    """延迟加载 OCR 模型，避免每次上传都重复初始化模型。"""
    return RapidOCR()


def _ocr_pdf_pages(content: bytes, page_indexes: list[int]) -> dict[int, str]:
    """把指定 PDF 页面渲染后离线识别；串行保护 PDFium 和 OCR 引擎。"""
    if len(page_indexes) > MAX_OCR_PAGES:
        raise DocumentParseError(f"扫描版 PDF 最多支持 {MAX_OCR_PAGES} 页 OCR")

    recognized: dict[int, str] = {}
    with _OCR_LOCK:
        pdf = pdfium.PdfDocument(content)
        try:
            for page_index in page_indexes:
                page = pdf[page_index]
                bitmap = None
                try:
                    bitmap = page.render(scale=OCR_DPI / 72)
                    result = _ocr_engine()(bitmap.to_pil())
                    texts = getattr(result, "txts", None) or []
                    recognized[page_index] = "\n".join(
                        str(text).strip() for text in texts if str(text).strip()
                    )
                finally:
                    if bitmap is not None:
                        bitmap.close()
                    page.close()
        finally:
            pdf.close()
    return recognized


def _decode_text(content: bytes, filename: str) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise DocumentParseError(f"{filename} 必须使用 UTF-8 编码") from error


def _parse_docx(content: bytes, filename: str) -> str:
    try:
        document = Document(BytesIO(content))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs]
        # Word 表格经常包含政策和参数，不能只读取普通段落。
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
        return "\n".join(line for line in lines if line)
    except Exception as error:
        raise DocumentParseError(f"{filename} 不是有效的 DOCX 文档") from error


def _parse_pdf(content: bytes, filename: str) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        # 只有完全没有文字层的页面才走 OCR，避免短标题页被重复识别。
        ocr_indexes = [index for index, text in enumerate(pages) if not text]
        if ocr_indexes:
            for index, text in _ocr_pdf_pages(content, ocr_indexes).items():
                if text:
                    pages[index] = text
        return "\n\n".join(page for page in pages if page)
    except DocumentParseError:
        raise
    except Exception as error:
        raise DocumentParseError(f"{filename} 不是有效或未加密的 PDF 文档") from error


def parse_document(filename: str, content: bytes) -> str:
    """根据扩展名安全解析文档正文，不执行文档中的宏、脚本或链接。"""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        if extension == ".doc":
            raise DocumentParseError("旧版 DOC 暂不支持，请另存为 DOCX 后上传")
        raise DocumentParseError("仅支持 TXT、Markdown、DOCX 和 PDF 文件")
    if extension in {".txt", ".md"}:
        return _decode_text(content, filename)
    if extension == ".docx":
        return _parse_docx(content, filename)
    return _parse_pdf(content, filename)
